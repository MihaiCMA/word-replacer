import os
from tkinter import Tk, Button, Label, filedialog, DISABLED, NORMAL
from docx import Document
from openpyxl import load_workbook

class WordReplacerApp:
    def __init__(self, master):
        self.master = master
        master.title("Word Replacer")
        master.geometry("400x150")

        self.label_template = Label(master, text="Select Word Template:")
        self.label_template.grid(row=0, column=0, padx=10, pady=10)

        self.template_button = Button(master, text="Choose Template", command=self.choose_template)
        self.template_button.grid(row=0, column=1, padx=10, pady=10)

        self.label_excel = Label(master, text="Select Excel File:")
        self.label_excel.grid(row=1, column=0, padx=10, pady=10)

        self.excel_button = Button(master, text="Choose Excel File", command=self.choose_excel, state=DISABLED)
        self.excel_button.grid(row=1, column=1, padx=10, pady=10)

        self.replace_button = Button(master, text="Replace Text", state=DISABLED, command=self.replace_text)
        self.replace_button.grid(row=2, column=1, padx=10, pady=10)

    def choose_template(self):
        self.template_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.template_path:
            self.excel_button.config(state=NORMAL)

    def choose_excel(self):
        self.excel_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if self.excel_path:
            self.replace_button.config(state=NORMAL)

    def replace_text(self):
        if hasattr(self, 'template_path') and hasattr(self, 'excel_path'):
            document = Document(self.template_path)
            wb = load_workbook(filename=self.excel_path)
            ws = wb.active
            output_folder = os.path.dirname(self.template_path)
            # Skip the first row (header)
            rows = iter(ws.iter_rows(values_only=True))
            next(rows)
            for i, row in enumerate(rows, start=1):
                old_text = "old_string"
                new_text = row[0]
                new_doc = Document(self.template_path)
                for paragraph in new_doc.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text.strip())
                output_filename = os.path.splitext(os.path.basename(self.template_path))[0] + f"_output_{i}.docx"
                output_path = os.path.join(output_folder, output_filename)
                new_doc.save(output_path)
            os.startfile(output_folder)
            self.master.destroy()
        else:
            print("Please select a Word template and Excel file first.")

root = Tk()
app = WordReplacerApp(root)
root.mainloop()
